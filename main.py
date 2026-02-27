import google.generativeai as genai
from groq import Groq
import pandas as pd
import json
import time
import re
import os
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pypandoc

# Load environment variables from .env file
load_dotenv()

# 1. Configuration
AI_PROVIDER = os.getenv("AI_PROVIDER", "gemini").lower()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Pipeline Control Flags
GENERATE_ANSWERS_AI = True       # Set to False to skip AI generation answers
GENERATE_SHEETS_FROM_EXCEL = True  # Set to True to extract CSVs from existing Excel if AI is skipped
GENERATE_MCQ_BOOK = True           # Set to True to generate DOCX/EPUB

# Batch Folder List
FOLDER_PATHS = [
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified AI Practitioner",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Advanced Networking – Specialty",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Cloud Practitioner",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified CloudOps Engineer – Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Data Engineer – Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified DevOps Engineer – Professional",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Developer – Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Machine Learning Engineer – Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Machine Learning – Specialty",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Security – Specialty",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Solutions Architect – Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\AWS Certified Solutions Architect – Professional",
    r"C:\git\generate-questions\question_generation\question_generation\output\AZ-800 Administering Windows Server Hybrid Core Infrastructure",
    r"C:\git\generate-questions\question_generation\question_generation\output\Associate Cloud Engineer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Associate Data Practitioner Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Associate Google Workspace Administrator Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Blue Prism Certified Developer",
    r"C:\git\generate-questions\question_generation\question_generation\output\CCNA",
    r"C:\git\generate-questions\question_generation\question_generation\output\CFA Program Level I",
    r"C:\git\generate-questions\question_generation\question_generation\output\CFA Program Level II",
    r"C:\git\generate-questions\question_generation\question_generation\output\CKA Certified Kubernetes Administrator",
    r"C:\git\generate-questions\question_generation\question_generation\output\CKAD Certified Kubernetes Application Developer",
    r"C:\git\generate-questions\question_generation\question_generation\output\CKS Certified Kubernetes Security Specialist",
    r"C:\git\generate-questions\question_generation\question_generation\output\COBIT® 2019 Foundation Certificate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certificate of Cloud Security Knowledge v.5",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certification of Capability in Business Analysis™",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified Application Developer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified Associate in Project Management (CAPM)®",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified Cloud Security Professional (CCSP)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified Information Security Manager® (CISM)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified Information Systems Auditor® (CISA)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified Information Systems Security Professional (CISSP)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified Internal Auditor (CIA)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified Management Accountant",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified SAFe® 6 Agilist",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified SAFe® 6 Product OwnerProduct Manager",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified SAFe® 6 Scrum Master",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified System Administrator",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified in Cybersecurity (CC)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified in Governance, Risk and Compliance (CGRC)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Certified in Risk and Information Systems Control™ (CRISC)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Check Point Certified Security Expert CCSE R81",
    r"C:\git\generate-questions\question_generation\question_generation\output\Cisco Certified CyberOps Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Cisco Certified Internetwork Expert Data Center (CCIE Data Center)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Cisco Certified Internetwork Expert Enterprise Infrastructure (CCIE Enterprise Infrastructure)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Cisco Certified Internetwork Expert Security (CCIE Security)",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA CySA+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA Data+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA DataSys+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA Linux+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA Network+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA PenTest+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA Project+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA Security+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA SecurityX Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA Server+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\CompTIA Tech+ Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\DAMA CDMP Certification - Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\DAMA CDMP Certification - Practitioner",
    r"C:\git\generate-questions\question_generation\question_generation\output\Databricks Certified Associate Developer for Apache Spark 3.0",
    r"C:\git\generate-questions\question_generation\question_generation\output\Databricks Certified Data Engineer Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\F5 Certified! Administrator, BIG-IP (F5-CA, BIG-IP)",
    r"C:\git\generate-questions\question_generation\question_generation\output\FE Exam",
    r"C:\git\generate-questions\question_generation\question_generation\output\Generative AI Leader Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\GitHub Actions",
    r"C:\git\generate-questions\question_generation\question_generation\output\GitHub Administration",
    r"C:\git\generate-questions\question_generation\question_generation\output\GitHub Advanced Security",
    r"C:\git\generate-questions\question_generation\question_generation\output\GitHub Copilot",
    r"C:\git\generate-questions\question_generation\question_generation\output\GitHub Foundations",
    r"C:\git\generate-questions\question_generation\question_generation\output\HashiCorp Certified Consul Associate (003)",
    r"C:\git\generate-questions\question_generation\question_generation\output\HashiCorp Certified Terraform Associate (004)",
    r"C:\git\generate-questions\question_generation\question_generation\output\HashiCorp Certified Vault Associate (003)",
    r"C:\git\generate-questions\question_generation\question_generation\output\ISTQB® Certified Tester - AI Testing (CT-AI)",
    r"C:\git\generate-questions\question_generation\question_generation\output\ISTQB® Certified Tester - Mobile Application Testing (CT-MAT)",
    r"C:\git\generate-questions\question_generation\question_generation\output\ISTQB® Certified Tester Advanced Level - Technical Test Analyst (CTAL-TTA)",
    r"C:\git\generate-questions\question_generation\question_generation\output\ISTQB® Certified Tester Advanced Level - Test Analyst (CTAL-TA)",
    r"C:\git\generate-questions\question_generation\question_generation\output\ISTQB® Certified Tester Advanced Level - Test Manager (CTAL-TM)",
    r"C:\git\generate-questions\question_generation\question_generation\output\ISTQB® Certified Tester Foundation Level (CTFL)",
    r"C:\git\generate-questions\question_generation\question_generation\output\ISTQB® Certified Tester Foundation Level - Agile Tester (CTFL-AT)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Juniper Networks Certified Associate, Junos (JNCIA-Junos)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Juniper Networks Certified Specialist, Service Provider Routing & Switching (JNCIS-SP)",
    r"C:\git\generate-questions\question_generation\question_generation\output\KCNA Kubernetes and Cloud Native Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Lean Six Sigma Black Belt Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Lean Six Sigma Green Belt Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Lean Six Sigma White Belt Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Lean Six Sigma Yellow Belt Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Linux Essentials Certificate",
    r"C:\git\generate-questions\question_generation\question_generation\output\MB-240 Microsoft Dynamics 365 Field Service Functional Consultant",
    r"C:\git\generate-questions\question_generation\question_generation\output\MB-500 Microsoft Dynamics 365 Finance and Operations Apps Developer",
    r"C:\git\generate-questions\question_generation\question_generation\output\MCE Microsoft Certified Educator",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft 365 Certified Fundamentals",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft 365 Certified Teams Administrator Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure AI Fundamentals",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Administrator Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Data Fundamentals",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Data Scientist Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Database Administrator Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Developer Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Fundamentals",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Network Engineer Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Security Engineer Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Solutions Architect Expert",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Azure Virtual Desktop Specialty",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Cybersecurity Architect Expert",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified DevOps Engineer Expert",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Dynamics 365 Field Service Functional Consultant Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Dynamics 365 Finance Functional Consultant Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Dynamics 365 Finance and Operations Apps Developer Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Dynamics 365 Finance and Operations Apps Solution Architect Expert",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Identity and Access Administrator Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Power BI Data Analyst Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Power Platform Developer Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Power Platform Functional Consultant Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Power Platform Fundamentals",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Security Operations Analyst Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Security, Compliance, and Identity Fundamentals",
    r"C:\git\generate-questions\question_generation\question_generation\output\Microsoft Certified Windows Server Hybrid Administrator Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\Oracle Database PLSQL Developer Certified Professional",
    r"C:\git\generate-questions\question_generation\question_generation\output\Oracle Database SQL Certified Associate",
    r"C:\git\generate-questions\question_generation\question_generation\output\PL-600 Microsoft Power Platform Solution Architect",
    r"C:\git\generate-questions\question_generation\question_generation\output\PMI Agile Certified Practitioner (PMI-ACP)®",
    r"C:\git\generate-questions\question_generation\question_generation\output\PMI Professional in Business Analysis (PMI-PBA)®",
    r"C:\git\generate-questions\question_generation\question_generation\output\PMI Risk Management Professional (PMI-RMP)®",
    r"C:\git\generate-questions\question_generation\question_generation\output\PTCB Certified Pharmacy Technician",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Cloud Architect",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Cloud Database Engineer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Cloud DevOps Engineer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Cloud Developer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Cloud Network Engineer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Cloud Security Engineer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Data Engineer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Machine Learning Engineer",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional Scrum Product Owner™ I (PSPO I)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Professional in Human Resources® (PHR®)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Program Management Professional (PgMP)®",
    r"C:\git\generate-questions\question_generation\question_generation\output\Project Management Professional (PMP)®",
    r"C:\git\generate-questions\question_generation\question_generation\output\Red Hat Certified Engineer (RHCE)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Red Hat Certified System Administrator (RHCSA)",
    r"C:\git\generate-questions\question_generation\question_generation\output\SAP Certified Associate - Back-End Developer - ABAP Cloud",
    r"C:\git\generate-questions\question_generation\question_generation\output\SAP Certified Associate - SAP Activate Project Manager",
    r"C:\git\generate-questions\question_generation\question_generation\output\SAP Certified Associate - SAP S4HANA Cloud Private Edition, Extended Warehouse Management",
    r"C:\git\generate-questions\question_generation\question_generation\output\SAP Certified Associate - SAP S4HANA Cloud Private Edition, Transportation Management",
    r"C:\git\generate-questions\question_generation\question_generation\output\SAP Certified Professional - SAP Enterprise Architect",
    r"C:\git\generate-questions\question_generation\question_generation\output\SAS Certified Professional Advanced Programming Using SAS 9.4",
    r"C:\git\generate-questions\question_generation\question_generation\output\SAS Certified Specialist Base Programming Using SAS 9.4",
    r"C:\git\generate-questions\question_generation\question_generation\output\SHRM Certified Professional (SHRM-CP)",
    r"C:\git\generate-questions\question_generation\question_generation\output\SHRM Senior Certified Professional (SHRM-SCP)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Six Sigma Black Belt Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Six Sigma Green Belt Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Six Sigma White Belt Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Six Sigma Yellow Belt Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\SnowPro Advanced Architect",
    r"C:\git\generate-questions\question_generation\question_generation\output\SnowPro Core Certification",
    r"C:\git\generate-questions\question_generation\question_generation\output\Splunk Core Certified Power User",
    r"C:\git\generate-questions\question_generation\question_generation\output\Splunk Core Certified User",
    r"C:\git\generate-questions\question_generation\question_generation\output\Splunk Enterprise Certified Admin",
    r"C:\git\generate-questions\question_generation\question_generation\output\Spring Certified Professional 2024 [v2]",
    r"C:\git\generate-questions\question_generation\question_generation\output\Systems Security Certified Practitioner (SSCP)",
    r"C:\git\generate-questions\question_generation\question_generation\output\Tableau Certified Data Analyst",
    r"C:\git\generate-questions\question_generation\question_generation\output\Tableau Desktop Specialist",
    r"C:\git\generate-questions\question_generation\question_generation\output\VMware Certified Professional - Data Center Virtualization",
    r"C:\git\generate-questions\question_generation\question_generation\output\VMware Certified Professional - Network Virtualization",
    r"C:\git\generate-questions\question_generation\question_generation\output\VMware Certified Technical Associate - Data Center Virtualization",
    r"C:\git\generate-questions\question_generation\question_generation\output\[PCAP-31-03] PCAP – Certified Associate Python Programmer",
    r"C:\git\generate-questions\question_generation\question_generation\output\[PCEP-30-02] PCEP™ – Certified Entry-Level Python Programmer",
]


INPUT_FILE = "questions.txt"

# Initialize AI Client
model = None
groq_client = None

if AI_PROVIDER == "gemini":
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not found in .env file")
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-2.5-flash')
    print("Using Gemini API for generation.")
elif AI_PROVIDER == "groq":
    if not GROQ_API_KEY:
        raise ValueError("GROQ_API_KEY not found in .env file")
    groq_client = Groq(api_key=GROQ_API_KEY)
    print("Using Groq API (Llama 3) for generation.")

columns = [
    "Question", "Question Type", 
    "Answer Option 1", "Explanation 1", 
    "Answer Option 2", "Explanation 2", 
    "Answer Option 3", "Explanation 3", 
    "Answer Option 4", "Explanation 4", 
    "Answer Option 5", "Explanation 5", 
    "Answer Option 6", "Explanation 6", 
    "Correct Answers", "Overall Explanation", "Domain"
]

def get_topic_name(folder_path):
    """Extract clean topic name from folder name."""
    folder_name = os.path.basename(folder_path.rstrip(os.sep))
    # Remove " Interview Questions Practice Test" if present or " Interview Questions"
    topic = folder_name.replace(" Interview Questions Practice Test", "").replace(" Interview Questions", "").strip()
    # Replace en-dash with standard hyphen to avoid encoding issues in Domain/Filenames
    topic = topic.replace("–", "-")
    return topic

def load_questions(folder_path, file_name):
    questions = []
    file_path = os.path.join(folder_path, file_name)
    if not os.path.exists(file_path):
        print(f"Error: {file_path} not found.")
        return []

    with open(file_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or re.match(r'^Section \d+.*', line, re.IGNORECASE):
                continue
            clean_q = re.sub(r'^(\d+[\.\)]|Question \d+:?)\s*', '', line, flags=re.IGNORECASE).strip()
            if clean_q:
                questions.append(clean_q)
    return questions

def get_ai_response(prompt):
    if AI_PROVIDER == "gemini":
        response = model.generate_content(prompt)
        return response.text.strip()
    elif AI_PROVIDER == "groq":
        chat_completion = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.1-8b-instant",
            response_format={"type": "json_object"}
        )
        return chat_completion.choices[0].message.content.strip()

def generate_question_data(question, total_count, current_index, topic_name, target_type):
    print(f"Processing question {current_index}/{total_count} ({target_type}): {question[:50]}...")
    
    type_instruction = ""
    if target_type == "multi-select":
        type_instruction = 'Ensure there are TWO or MORE correct answers. "Question Type" MUST be "multi-select". "Correct Answers" MUST be a comma-separated list of 2 or more digits (e.g. "1, 3").'
    else:
        type_instruction = 'Ensure there is exactly ONE correct answer. "Question Type" MUST be "multiple-choice". "Correct Answers" MUST be a single digit (e.g. "1").'

    prompt = f"""
    You are an expert {topic_name} coach. 
    Analyze the following {topic_name} interview question and generate 6 multiple-choice options, explanations for each, the correct answer indices, an overall explanation, and the domain.
    
    Question: "{question}"
    
    Respond ONLY with a valid JSON object matching this exact structure, with no markdown formatting or extra text.
    
    IMPORTANT: {type_instruction}
    
    The "Question" in the JSON must NOT contain any numbering prefix.
    
    {{
      "Question": "The clean question text",
      "Question Type": "{target_type}",
      "Answer Option 1": "Option 1 text",
      "Explanation 1": "Explanation why 1 is right/wrong",
      "Answer Option 2": "Option 2 text",
      "Explanation 2": "Explanation why 2 is right/wrong",
      "Answer Option 3": "Option 3 text",
      "Explanation 3": "Explanation why 3 is right/wrong",
      "Answer Option 4": "Option 4 text",
      "Explanation 4": "Explanation why 4 is right/wrong",
      "Answer Option 5": "Option 5 text",
      "Explanation 5": "Explanation why 5 is right/wrong",
      "Answer Option 6": "Option 6 text",
      "Explanation 6": "Explanation why 6 is right/wrong",
      "Correct Answers": "1",
      "Overall Explanation": "A comprehensive summary of the concept",
      "Domain": "{topic_name} topic area"
    }}
    """
    
    attempt = 1
    while True:
        try:
            response_text = get_ai_response(prompt)
            if response_text.startswith("```json"):
                response_text = response_text[7:-3].strip()
            elif response_text.startswith("```"):
                response_text = response_text[3:-3].strip()
            
            data = json.loads(response_text)
            
            # Use the AI's provided Question Type but ensure Correct Answers matches the requested logic
            correct_ans_str = str(data.get("Correct Answers", ""))
            correct_ans_list = [c.strip() for c in re.split(r'[;,]', correct_ans_str) if c.strip()]
            
            # Final validation/enforcement based on target_type
            if target_type == "multi-select" and len(correct_ans_list) < 2:
                # If AI failed to provide multiple answers, try to force it via retry logic or just let it pass if critical
                # For now, we trust the AI instruction is strong enough, but we update the Type to reflect truth
                data["Question Type"] = "multiple-choice" if len(correct_ans_list) <= 1 else "multi-select"
            elif target_type == "multiple-choice" and len(correct_ans_list) > 1:
                data["Question Type"] = "multi-select"
            
            # Normalize single answer
            if len(correct_ans_list) == 1:
                data["Correct Answers"] = correct_ans_list[0]

            # Ensure required keys exist to avoid downstream errors
            if all(key in data for key in columns):
                return data
            else:
                print(f"Incomplete JSON received for question {current_index}. Retrying...")
        except Exception as e:
            wait_time = min(2 ** attempt, 30) # Exponential backoff capped at 30s
            print(f"Error for question {current_index} (Attempt {attempt}): {e}. Retrying in {wait_time}s...")
            time.sleep(wait_time)
            attempt += 1

def clean_all_text(text):
    if not isinstance(text, str): return text
    patterns = ["A. ", "B. ", "C. ", "D. ", "E. ", "F. ", "1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "a. ", "b. ", "c. ", "d. ", "e. ", "f. "]
    for p in patterns: text = text.replace(p, "")
    return text.strip()

def format_correct_answers(ans_val):
    """Ensure correct answers are comma-separated with no spaces (e.g., '1,2,3')."""
    if pd.isna(ans_val): return ""
    ans_str = str(ans_val)
    # Remove all spaces and ensure standard comma delimiter
    ans_str = ans_str.replace(" ", "")
    # If it was something like '1;2', clean that too
    ans_str = ans_str.replace(";", ",")
    return ans_str

def merge_csv_files(folder_path, output_file):
    csv_files = [file for file in os.listdir(folder_path) if file.endswith(".csv") and file.startswith("Sheet")]
    if not csv_files: return None
    dfs = []
    for file in csv_files:
        try:
            df = pd.read_csv(os.path.join(folder_path, file))
            dfs.append(df)
        except Exception as e: print(f"Error reading {file}: {e}")
    if not dfs: return None
    merged_df = pd.concat(dfs, ignore_index=True)
    columns_to_keep = ["Question", "Question Type", "Answer Option 1", "Explanation 1", "Answer Option 2", "Explanation 2", "Answer Option 3", "Explanation 3", "Answer Option 4", "Explanation 4", "Answer Option 5", "Explanation 5", "Answer Option 6", "Explanation 6", "Correct Answers", "Overall Explanation", "Domain"]
    existing_columns = [c for c in columns_to_keep if c in merged_df.columns]
    merged_df = merged_df[existing_columns].dropna(subset=["Question", "Correct Answers", "Overall Explanation"])
    merged_df["Correct Answers"] = merged_df["Correct Answers"].apply(format_correct_answers)
    merged_df.to_csv(output_file, index=False, encoding='utf-8-sig')
    return merged_df

def create_docx(df, title_name, output_file):
    doc = Document()
    title = doc.add_heading(title_name, 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Exam Prep and Study Guide\n"); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; sub.runs[0].bold = True
    total_text = doc.add_paragraph(f"Total Questions: {len(df)}"); total_text.alignment = WD_ALIGN_PARAGRAPH.CENTER; total_text.runs[0].italic = True
    author = doc.add_paragraph("\nBy\nManish Dnyandeo Salunke"); author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading("Preface", level=1)
    content = (
        "In today's fast-paced and highly competitive tech world, interviews have evolved beyond just technical know-how. "
        "They now require clarity, speed, and confidence in answering structured questions.\n\n"
        "This book is a practical resource for aspiring system engineers, students, and IT professionals preparing for technical interviews. "
        "It covers a wide range of multiple-choice questions (MCQs), complete with correct answers and concise explanations to help reinforce your understanding.\n\n"
        "Whether you're preparing for your first job, transitioning into a new role, or simply brushing up on your skills, "
        "this book is designed to serve as a quick and effective learning tool.\n\n"
        "Thank you for choosing this book as part of your preparation journey. I hope it helps you succeed and grow in your IT career.\n\n"
        "— Manish Dnyandeo Salunke"
    )
    doc.add_paragraph(content)
    doc.add_page_break()
    doc.add_heading("About the Author", level=1)
    bio = (
        "Manish Dnyandeo Salunke is a seasoned IT professional, educator, and passionate author from Pune, India. "
        "With years of hands-on experience in the IT industry, Manish has contributed to various roles involving system engineering, infrastructure management, and technical support.\n\n"
        "His passion for writing and mentoring led him to create practical learning resources aimed at helping aspiring IT professionals succeed in their careers.\n\n"
        "Outside his technical pursuits, Manish enjoys storytelling, content creation, and writing books that simplify complex concepts for everyone."
    )
    doc.add_paragraph(bio)
    doc.add_page_break()
    KEYS = [("Answer Option 1", "Explanation 1"), ("Answer Option 2", "Explanation 2"), ("Answer Option 3", "Explanation 3"), ("Answer Option 4", "Explanation 4"), ("Answer Option 5", "Explanation 5"), ("Answer Option 6", "Explanation 6")]
    LABELS = ["A", "B", "C", "D", "E", "F"]
    for q_num, (_, row) in enumerate(df.iterrows(), start=1):
        doc.add_heading(f"Q{q_num}. {row['Question']}", level=1)
        if "Domain" in row: doc.add_paragraph(f"Domain: {row['Domain']}").runs[0].italic = True
        doc.add_paragraph()

        options_data = []
        for i, (opt_col, exp_col) in enumerate(KEYS):
            if opt_col in row and not pd.isna(row[opt_col]):
                opt_text = str(row[opt_col])
                exp_text = str(row[exp_col]) if exp_col in row and not pd.isna(row[exp_col]) else ""
                options_data.append((LABELS[i], opt_text, exp_text))
                p = doc.add_paragraph()
                p.add_run(f"{LABELS[i]}. ").bold = True
                p.add_run(opt_text)
        
        doc.add_paragraph()

        # Parse correct answers
        correct_ans_raw = str(row["Correct Answers"])
        correct_labels = []
        for part in re.split(r'[;,]', correct_ans_raw):
            part = part.strip().upper()
            if part.isdigit():
                idx = int(part) - 1
                if 0 <= idx < len(LABELS):
                    correct_labels.append(LABELS[idx])
            elif part in LABELS:
                correct_labels.append(part)
        
        display_ans = ", ".join(correct_labels) if correct_labels else correct_ans_raw
        
        ca_p = doc.add_paragraph()
        ca_p.add_run("Correct Answer: ").bold = True
        ca_p.add_run(display_ans)

        if "Overall Explanation" in row:
            exp_p = doc.add_paragraph()
            exp_p.add_run("Explanation: ").bold = True
            exp_p.add_run(str(row["Overall Explanation"]))

        # Answer Analysis Section
        if any(exp for _, _, exp in options_data if exp):
            doc.add_paragraph()
            hdr_p = doc.add_paragraph()
            hdr_p.add_run("Answer Analysis:").bold = True

            for label, opt_text, exp_text in options_data:
                if not exp_text: continue
                is_correct = label in correct_labels
                analysis_p = doc.add_paragraph()
                status = "Correct" if is_correct else "Wrong"
                analysis_p.add_run(f"{label}. [{status}] ").bold = True
                analysis_p.add_run(f"{opt_text}: ").italic = True
                analysis_p.add_run(exp_text)

        doc.add_paragraph()
        doc.add_page_break()
    doc.add_heading("Copyright Disclaimer", level=1)
    text = (
        "© Manish Dnyandeo Salunke. All rights reserved.\n\n"
        "No part of this book may be reproduced or transmitted in any form or by any means—electronic, mechanical, "
        "photocopying, recording, or otherwise—without the prior written permission of the author, "
        "except for brief quotations used in reviews or educational contexts.\n\n"
        "For permissions, please contact the author directly."
    )
    doc.add_paragraph(text)
    doc.save(output_file)

def convert_docx_to_epub(docx_file, epub_file, title):
    try:
        pypandoc.convert_file(docx_file, 'epub', outputfile=epub_file, extra_args=[f"--metadata=title:{title}", "--metadata=author:Manish Dnyandeo Salunke", "--metadata=lang:en"])
    except Exception as e: print(f"EPUB error: {e}")

def process_single_folder(folder_path):
    topic_name = get_topic_name(folder_path)
    print(f"\n========================================")
    print(f"PROCESSING TOPIC: {topic_name}")
    print(f"FOLDER: {folder_path}")
    print(f"========================================\n")
    
    if not os.path.exists(folder_path):
        print(f"Error: Folder does not exist: {folder_path}")
        return
        
    # AI Generation Phase
    if GENERATE_ANSWERS_AI:
        questions = load_questions(folder_path, INPUT_FILE)
        if not questions: return
        
        # Truncate topic name for filename to avoid Windows MAX_PATH (260 chars) issues
        safe_topic = topic_name.replace(' ', '_')
        if len(safe_topic) > 50:
            safe_topic = safe_topic[:50]
        
        output_xlsx = f"{safe_topic}_Generated.xlsx"
        output_path = os.path.join(folder_path, output_xlsx)
        
        writer = pd.ExcelWriter(output_path, engine='xlsxwriter')
        
        chunk_size = 250
        for i in range(0, len(questions), chunk_size):
            chunk = questions[i:i + chunk_size]
            chunk_index = (i // chunk_size) + 1
            print(f"\n--- Sheet {chunk_index} ({i+1} to {i+len(chunk)}) ---")
            chunk_rows = []
            
            # Calculate how many multi-select questions we need in this chunk (20%)
            multi_select_count = max(1, int(len(chunk) * 0.20)) if len(chunk) >= 5 else (1 if len(chunk) > 0 else 0)
            multi_select_indices = [idx for idx in range(len(chunk)) if (idx + 1) % 5 == 0]
            if len(multi_select_indices) < multi_select_count:
                for idx in range(len(chunk)):
                    if idx not in multi_select_indices:
                        multi_select_indices.append(idx)
                        if len(multi_select_indices) >= multi_select_count: break
            
            for j, question in enumerate(chunk):
                target_type = "multi-select" if j in multi_select_indices else "multiple-choice"
                chunk_rows.append(generate_question_data(question, len(questions), i + j + 1, topic_name, target_type))
                time.sleep(2 if AI_PROVIDER == "gemini" else 1)
                if (j + 1) % 10 == 0 and (j + 1) < len(chunk):
                    print("--- Rate limit pause (5s) ---"); time.sleep(3)
            
            df = pd.DataFrame(chunk_rows, columns=columns)
            for col in df.columns: df[col] = df[col].apply(clean_all_text)
            # Apply correct answers formatting
            df["Correct Answers"] = df["Correct Answers"].apply(format_correct_answers)
            
            df.to_excel(writer, sheet_name=f'Sheet{chunk_index}', index=False)
            df.to_csv(os.path.join(folder_path, f"Sheet{chunk_index}.csv"), index=False, encoding='utf-8-sig')
        
        writer.close()
    
    # Excel to CSV Extraction Phase (only if AI generation was skipped)
    elif GENERATE_SHEETS_FROM_EXCEL:
        print("Skipping AI generation. Attempting to extract sheets from existing Excel file...")
        target_xlsx = None
        xlsx_files = [f for f in os.listdir(folder_path) if f.endswith("_Generated.xlsx")]
        if xlsx_files:
            target_xlsx = os.path.join(folder_path, xlsx_files[0])
        else:
            q_xlsx = os.path.join(folder_path, "questions.xlsx")
            if os.path.exists(q_xlsx):
                target_xlsx = q_xlsx
        
        if target_xlsx:
            print(f"Found Excel: {target_xlsx}")
            xls = pd.ExcelFile(target_xlsx)
            for i, sheet_name in enumerate(xls.sheet_names, start=1):
                df = pd.read_excel(xls, sheet_name=sheet_name)
                # Verify and format while generating CSV
                if "Correct Answers" in df.columns:
                    df["Correct Answers"] = df["Correct Answers"].apply(format_correct_answers)
                df.to_csv(os.path.join(folder_path, f"Sheet{i}.csv"), index=False, encoding='utf-8-sig')
                print(f"Extracted {sheet_name} to Sheet{i}.csv")
        else:
            print("No suitable Excel file found (*_Generated.xlsx or questions.xlsx).")

    # Book Generation Phase
    if GENERATE_MCQ_BOOK:
        merged_csv = os.path.join(folder_path, "Merged_Questions.csv")
        merged_df = merge_csv_files(folder_path, merged_csv)
        if merged_df is not None:
            title = os.path.basename(folder_path.rstrip(os.sep))
            docx_p = os.path.join(folder_path, "MCQ_Ebook.docx")
            epub_p = os.path.join(folder_path, "MCQ_Ebook.epub")
            create_docx(merged_df, title, docx_p)
            convert_docx_to_epub(docx_p, epub_p, title)
            print(f"Ebook generated: {docx_p}")
    
    print(f"\nFinished processing: {topic_name}")

def main():
    while FOLDER_PATHS:
        folder = FOLDER_PATHS[0]
        try:
            process_single_folder(folder)
        except Exception as e:
            print(f"CRITICAL ERROR in {folder}: {e}")
        
        # Remove folder from list once done
        FOLDER_PATHS.pop(0)
        
        if FOLDER_PATHS:
            print(f"\n--- Waiting 5 Secounds before next folder ({len(FOLDER_PATHS)} folders remaining) ---")
            time.sleep(5)

if __name__ == "__main__":
    main()
